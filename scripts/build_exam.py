#!/usr/bin/env python3
"""
build_exam.py — Populate an exam cover sheet with metadata from _metadata.yml
                and merge with a rendered exam docx.

Usage
-----
    python scripts/build_exam.py <exam.qmd|exam.docx> [options]

Options
-------
    --cover   in-person | online   (default: in-person)
    --output  path/to/output.docx  (default: <input-stem>-final.docx next to input)

Examples
--------
    python scripts/build_exam.py exams/class-test/class-test-lectures-1-5.qmd
    python scripts/build_exam.py exams/class-test/class-test-lectures-1-5.docx --cover online
    python scripts/build_exam.py exams/class-test/class-test-lectures-1-5.qmd --output exam-print.docx
"""

import argparse
import subprocess
import sys
import yaml
from pathlib import Path

from docx import Document
from docxcompose.composer import Composer

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

PROJECT_ROOT = Path(__file__).resolve().parent.parent
METADATA_FILE = PROJECT_ROOT / "exams" / "_metadata.yml"

COVER_SHEETS = {
    "in-person": PROJECT_ROOT / "exams" / "In-person exam paper cover sheet_2025-26_S2.docx",
    "online":    PROJECT_ROOT / "exams" / "Online exam paper cover sheet_2025-26_S2.docx",
}

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def replace_in_para(para, old: str, new: str) -> bool:
    """Replace *old* with *new* inside *para*, collapsing runs as needed.

    Runs are collapsed into the first run so that text split across multiple
    runs (e.g. '[', 'content', ']') is handled correctly.  The first run's
    character formatting is preserved; subsequent run text is cleared.
    """
    if old not in para.text:
        return False
    new_text = para.text.replace(old, new)
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    return True


def replace_all(doc, old: str, new: str) -> None:
    """Apply replace_in_para across every paragraph in *doc*."""
    for para in doc.paragraphs:
        replace_in_para(para, old, new)


def replace_matching(doc, context: str, old: str, new: str) -> None:
    """Replace *old* → *new* only in the first paragraph that contains *context*."""
    for para in doc.paragraphs:
        if context in para.text and old in para.text:
            replace_in_para(para, old, new)
            return


def replace_c_instruction(doc, new_text: str) -> None:
    """Replace the entire (c) instruction paragraph."""
    for para in doc.paragraphs:
        if "(c) [" in para.text or ("(c)" in para.text and "Please state" in para.text):
            if para.runs:
                para.runs[0].text = f"(c) {new_text}"
                for run in para.runs[1:]:
                    run.text = ""
            return


# ---------------------------------------------------------------------------
# Main logic
# ---------------------------------------------------------------------------

def populate_cover(cover_path: Path, meta: dict) -> Document:
    """Open *cover_path* and fill in all metadata placeholders."""
    doc = Document(str(cover_path))
    instr = meta.get("exam-instructions", {})

    # Simple token replacements (order matters: longer tokens first)
    replace_matching(doc, "MODULE CODE",   "XXyyyy", meta.get("module-code", ""))
    replace_matching(doc, "MODULE TITLE",  "xxxx",   meta.get("module-title", ""))
    replace_matching(doc, "EXAM DURATION", "xxxx",   meta.get("exam-duration", ""))

    # Instruction (a)
    replace_in_para_globally(doc, "[Number of questions to be answered]", instr.get("a", ""))
    # Instruction (b)
    replace_in_para_globally(doc, "[Weighting of questions/sections]",    instr.get("b", ""))
    # Instruction (c) — contains curly-quoted characters; match by context
    replace_c_instruction(doc, instr.get("c", ""))
    # Instruction (d)
    instr_d = instr.get("d") or "None."
    replace_in_para_globally(doc, "[Any other instructions]", instr_d)

    return doc


def replace_in_para_globally(doc, old: str, new: str) -> None:
    """Apply replace_in_para to every paragraph containing *old*."""
    for para in doc.paragraphs:
        replace_in_para(para, old, new)


def find_rendered_docx(qmd_path: Path) -> Path:
    """Locate the docx that Quarto produces for *qmd_path*."""
    # Quarto project output goes to _site/
    try:
        rel = qmd_path.relative_to(PROJECT_ROOT)
        candidate = PROJECT_ROOT / "_site" / rel.with_suffix(".docx")
        if candidate.exists():
            return candidate
    except ValueError:
        pass
    # Fallback: same directory
    return qmd_path.with_suffix(".docx")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Build exam paper: populate cover sheet and merge with exam content."
    )
    parser.add_argument("input", help="Path to .qmd or .docx exam file")
    parser.add_argument(
        "--cover", choices=["in-person", "online"], default="in-person",
        help="Cover sheet type (default: in-person)"
    )
    parser.add_argument("--output", help="Output docx path")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()

    # ── 1. Render QMD → docx if needed ─────────────────────────────────────
    if input_path.suffix == ".qmd":
        print(f"Rendering {input_path.name} …")
        subprocess.run(
            ["quarto", "render", str(input_path), "--to", "docx"],
            check=True,
        )
        docx_path = find_rendered_docx(input_path)
    else:
        docx_path = input_path

    if not docx_path.exists():
        sys.exit(f"Error: rendered docx not found at {docx_path}")

    # ── 2. Load metadata ────────────────────────────────────────────────────
    if not METADATA_FILE.exists():
        sys.exit(f"Error: metadata file not found at {METADATA_FILE}")
    with open(METADATA_FILE, encoding="utf-8") as f:
        meta = yaml.safe_load(f)

    # ── 3. Populate cover sheet ─────────────────────────────────────────────
    cover_path = COVER_SHEETS[args.cover]
    if not cover_path.exists():
        sys.exit(f"Error: cover sheet not found at {cover_path}")

    print(f"Populating cover sheet: {cover_path.name}")
    cover_doc = populate_cover(cover_path, meta)

    # ── 4. Merge cover + exam body ──────────────────────────────────────────
    print(f"Merging with exam body: {docx_path.name}")
    exam_doc = Document(str(docx_path))
    composer = Composer(cover_doc)
    composer.append(exam_doc)

    # ── 5. Save ─────────────────────────────────────────────────────────────
    if args.output:
        output_path = Path(args.output).resolve()
    else:
        output_path = docx_path.parent / (docx_path.stem + "-final.docx")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(output_path))
    print(f"Done: {output_path}")


if __name__ == "__main__":
    main()
